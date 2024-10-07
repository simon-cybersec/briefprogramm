import customtkinter
from pathlib import Path
# For creation of data directory
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Mm

import subprocess

customtkinter.set_appearance_mode("Light")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

#customtkinter.set_widget_scaling(1.5)  # widget dimensions and text size
#customtkinter.set_window_scaling(1.5)  # window geometry dimensions

class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # Configure window
        self.title("")
        self.filename = "Brief"
        self.path = self.create_brief_directory()
        
        self.geometry("1500x1000")       # width, height

        # Fullscreen
        self.wm_attributes('-fullscreen', True)
        self.state('normal')  # This call is appears to be necessary to make the app actually go full screen.

        custom_font_textbox = ("Times",38,'bold')
        custom_font_title = ("Times",42,'bold')
        custom_font_button = ("Times",24,'bold')

        # Padding values
        self.padding_x = 40
        self.padding_y = 20
        
        # Grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=0)
        self.grid_columnconfigure(2, weight=0)
        self.grid_columnconfigure(3, weight=1)

        self.grid_rowconfigure(0, weight=0)     # Title
        self.grid_rowconfigure(1, weight=1)     # Fertig und Drucken button
        self.grid_rowconfigure(2, weight=0)     # Textbox

        

        # Title
        self.label = customtkinter.CTkLabel(self,
                                            text="Brief Programm",
                                            font=custom_font_title)
        self.label.grid(row=0, column=0, columnspan=4, sticky="nsew", padx=self.padding_x, pady=self.padding_y)


        # Textbox
        self.textbox = customtkinter.CTkTextbox(self,
                                                corner_radius=8, 
                                                font=custom_font_textbox,
                                                border_width=2,
                                                border_color="black")
        self.textbox.focus_set()
        self.textbox.grid(row=1, column=0, columnspan=4, sticky="nsew", padx=self.padding_x +50, pady=self.padding_y)
        #self.textbox.insert("0.0", "Some example text!\n" * 10)


        # Button fertig
        self.button_fertig = customtkinter.CTkButton(self, 
                                                     text="Fertig", 
                                                     font=custom_font_button,
                                                     command=self.button_fertig_click)
        self.button_fertig.grid(row=2, column=1, padx=10, pady=self.padding_y, sticky="nwse")

        # Button drucken
        #self.button_drucken = customtkinter.CTkButton(self, 
        #                                             text="Drucken", 
        #                                             font=custom_font_button,
        #                                             command=self.button_fertig_click)
        #self.button_drucken.grid(row=2, column=2, padx=10, pady=self.padding_y, sticky="nwse")

        



    def button_fertig_click(self):

        text = self.textbox.get("0.0", "end")  # get text from line 0 character 0 till the end

        if text != " ":
            self.filename = "Brief_" + datetime.now().strftime('%Y-%m-%d_%H-%M-%S') + ".docx"

            document = Document()

            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(25.4)
            section.right_margin = Mm(25.4)
            section.top_margin = Mm(25.4)
            section.bottom_margin = Mm(25.4)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)

            style = document.styles['Normal']
            font = style.font
            font.name = 'Times'
            font.size = Pt(24)

            p = document.add_paragraph('')
            p.style = document.styles['Normal']
            p.add_run(text).bold = True

            document.save(self.path + self.filename)

            try:
                #libreoffice --headless --convert-to pdf filename.doc
                subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", self.path, self.path+self.filename]) 
                #subprocess.run(["abiword", "--to=pdf", self.filename]) 
            except:
                print("Sth went problematic when converting to pdf.") 

        # Exit program
        print("Exiting. File: ", self.filename)
        self.destroy() 


    def create_brief_directory(self):

        # Check if data directory structure already exists
        if not os.path.exists("./Briefe_von_Alex/"):
            os.mkdir("./Briefe_von_Alex/")
        
        # Create folder name
        path = "./Briefe_von_Alex/"

        # Create folder
        try:
            os.mkdir(path)
            print("Folder %s created. \n" % path)
        except FileExistsError:
            print("Folder %s exists." % path)

        return path



if __name__ == "__main__":
    app = App()
    app.mainloop()
