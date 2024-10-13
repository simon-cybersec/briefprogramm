import customtkinter
from tkinter.messagebox import askyesno

import os, sys

from docx import Document
from docx.shared import Pt, Mm

import subprocess
import csv

import contextlib

customtkinter.set_appearance_mode("Light")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

#customtkinter.set_widget_scaling(1.5)  # widget dimensions and text size
#customtkinter.set_window_scaling(1.5)  # window geometry dimensions

class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()        

        self.bind("<Escape>", self.close_briefprogramm)

        # Configure window
        self.title("")
                
        # Fullscreen
        self.geometry("1500x1000")       # width, height
        self.wm_attributes('-fullscreen', True)
        self.state('normal')  # This call is appears to be necessary to make the app actually go full screen.

        # Fonts 
        custom_font_textbox = ("Times",38,'bold')
        custom_font_title = ("Times",42,'bold')
        custom_font_button = ("Times",24,'bold')

        # Padding values
        self.padding_x = 40
        self.padding_y = 20
        
        # Grid
        # 4 Colums
        self.grid_columnconfigure(0, weight=1)
        self.grid_columnconfigure(1, weight=0)
        self.grid_columnconfigure(2, weight=0)
        self.grid_columnconfigure(3, weight=1)
        # 3 Rows
        self.grid_rowconfigure(0, weight=0)     # Title
        self.grid_rowconfigure(1, weight=1)     # Textbox
        self.grid_rowconfigure(2, weight=0)     # Fertig und Drucken button

        

        # Title
        self.label = customtkinter.CTkLabel(self,
                                            text="Brief Programm",
                                            font=custom_font_title)
        self.label.grid(row=0, column=0, columnspan=4, sticky="nsew", padx=self.padding_x, pady=self.padding_y)


        # Textbox
        self.textbox = customtkinter.CTkTextbox(self,
                                                corner_radius=0, 
                                                font=custom_font_textbox,
                                                border_width=2,
                                                border_color="black") #activate_scrollbars=False
                                                
        self.textbox.focus_set()
        self.textbox.grid(row=1, column=0, columnspan=4, sticky="nsew", padx=self.padding_x +80, pady=self.padding_y)
        #self.textbox.insert("0.0", "Some example text!\n" * 10)


        # Button fertig
        self.button_fertig = customtkinter.CTkButton(self, 
                                                     text="Fertig", 
                                                     font=custom_font_button,
                                                     command=self.button_fertig_click)
        self.button_fertig.grid(row=2, column=1, padx=10, pady=self.padding_y, sticky="nwse")


        ## create CTk scrollbar
        #textbox_scrollbar = customtkinter.CTkScrollbar(self, command=self.textbox.yview)
        #textbox_scrollbar.grid(row=1, column=3, sticky="ens")
        ## connect textbox scroll event to CTk scrollbar
        #self.textbox.configure(yscrollcommand=textbox_scrollbar.set)

        ########

        # Temporary filename
        self.filename = "temp_filename"

        # Directory to store the files at
        self.store_dir = "./Briefe/"
        self.create_storage_directory(self.store_dir)

        # Get Wifi SSID
        self.wifi_ssid = sys.argv[1]
        print("my wifi_ssid: " + self.wifi_ssid)
        
        # Get printer name:
        # The printer is then chosen depending on the wifi the computer is connected to.
        # In the myPrinters.csv is the mapping wifi-ssid -> printer name
        # The user needs to write it into the myPrinters.csv like: wifi_ssid, printer_name
        # Incase you need the name of your printer, go to http://localhost:631/printers (Cups need to be installed) and look for your printers name
        csv_filename = "myPrinters.csv"
        with open(csv_filename, newline='') as csvfile:
            printerreader = csv.reader(csvfile, delimiter=',')
            for row in printerreader:
                if row[0] == self.wifi_ssid:
                    self.printer_used = row[1]
                    print("Using this printer: " + self.printer_used)
                    break


    #
    # Clicking Button "Fertig"
    #
    def button_fertig_click(self):

        # Get written text
        text = self.textbox.get("0.0", "end")  # get text from line 0 character 0 till the end
        if len(text) > 1:

            # Switch mouse cursor to watch symbol
            with self.WaitCursor():

                # Popup asking for printing file
                answer = self.popup_print_yesno() # self.popup_print()

                # Clear textbox
                self.textbox.delete(0.0, 'end')

                # Get number of files in storage directory for file naming.
                #number_of_files = len(os.listdir(self.store_dir))
                number_of_briefe = len([b for b in os.listdir(self.store_dir) if b.endswith('.docx')])

                # Filename 
                self.filename = "Brief-" + str(number_of_briefe + 1)

                # Setting up a Din A4 document
                document = Document()
                section = document.sections[0]
                section.page_height     = Mm(297)
                section.page_width      = Mm(210)
                section.left_margin     = Mm(25.4)
                section.right_margin    = Mm(25.4)
                section.top_margin      = Mm(25.4)
                section.bottom_margin   = Mm(25.4)
                section.header_distance = Mm(12.7)
                section.footer_distance = Mm(12.7)

                style = document.styles['Normal']
                font = style.font
                font.name = 'Times'
                font.size = Pt(24)

                p = document.add_paragraph('')
                p.style = document.styles['Normal']
                p.add_run(text).bold = True

                # Save as .docx file
                document.save(self.store_dir + self.filename + ".docx")

                # Convert to pdf file
                self.convert_docx_to_pdf(docx_file=self.store_dir+self.filename+".docx", output_directory=self.store_dir)

                if answer == True:
                    self.print_pdf(printer_name=self.printer_used, pdf_file=self.store_dir+self.filename+".pdf")

                print("\nFile: " + self.filename + "\nExit.")
                
        else:
            print("Nothing written.")

        # Exit program     
        self.close_briefprogramm(event=None)


    #
    # Create storage directory for the files
    #
    def create_storage_directory(self, storage_directory):
        
        try:
            os.mkdir(storage_directory)
            print("Folder " + storage_directory + " created.")
        except FileExistsError:
            print("Folder " + storage_directory + "found.")

    
    #
    # Convert a docx file to pdf. (Libre office needs to be installed!)
    #
    def convert_docx_to_pdf(self, docx_file, output_directory):
        try:
            # Command: soffice --headless --convert-to pdf filename.doc
            subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_directory, docx_file]) 
            print("Converted successfully to PDF.")
        except:
            print("Sth went wrong when converting to PDF.") 

    #
    # Print PDF file
    #
    def print_pdf(self, printer_name, pdf_file):
        
        # For example: lp -d EPSON_ET_2850_Series Brief-1.pdf
        try:
            subprocess.run(["lp", "-d", printer_name, pdf_file])
            print("Printing...")
        except:
            print("Failed when trying to print.")
    
    
    #
    # Close window
    #
    def close_briefprogramm(self, event):
        self.destroy()


    #
    # Popup asking for printing the file
    #
    def popup_print_yesno(self):
        answer = askyesno(title=' ',
            message='BRIEF DRUCKEN ?')
        print("myanswer: " + str(answer))

        return answer
    

    #
    # New file
    #
    def new_file(self):
        self.filename = "temp_filename"
        self.textbox.delete(0.0, 'end')
        self.textbox.focus_set()
   

    #
    # Switch mouse cursor to watch symbol
    #
    @contextlib.contextmanager
    def WaitCursor(self):
        self.config(cursor="watch")
        self.update()
        try:
            yield self
        finally:
            self.config(cursor="")


# --- Program starts here ---
if __name__ == "__main__":
    app = App()
    app.mainloop()
